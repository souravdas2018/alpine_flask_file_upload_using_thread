# from quart import Quart, request, jsonify, render_template
# import aiofiles
# import os, uuid, asyncio, functools
# from io import BytesIO
# from docx import Document
# import PyPDF2
# import openai

# app = Quart(__name__)
# UPLOAD_FOLDER = "uploads"
# os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# tasks = {}
# extracted_contents = {}

# openai.api_key = os.getenv('OPENAI_API_KEY')  # Load from environment variable


# def process_document(file_contents):
#     text = ''
#     doc = Document(BytesIO(file_contents))
#     for paragraph in doc.paragraphs:
#         text += paragraph.text + '\n'
#     return text


# def process_pdf(file_contents):
#     text = ''
#     reader = PyPDF2.PdfReader(BytesIO(file_contents))
#     for page in reader.pages:
#         text += page.extract_text()
#     return text


# async def async_save_file(task_id, filename, content):
#     tasks[task_id]["files"][filename] = {"status": "processing", "done": False}
#     await asyncio.sleep(10)  # simulate processing delay

#     filepath = os.path.join(UPLOAD_FOLDER, filename)
#     async with aiofiles.open(filepath, "wb") as f:
#         await f.write(content)

#     tasks[task_id]["files"][filename] = {"status": "completed", "done": True}
#     if all(file["done"] for file in tasks[task_id]["files"].values()):
#         tasks[task_id]["status"] = "completed"


# @app.route("/")
# async def index():
#     return await render_template("index.html")


# @app.route("/upload", methods=["POST"])
# async def upload():
#     files = (await request.files).getlist("files")
#     if not files:
#         return jsonify({"error": "No files uploaded"}), 400

#     task_id = str(uuid.uuid4())
#     tasks[task_id] = {"status": "processing", "files": {}}

#     for f in files:
#         filename = f.filename
#         content = f.read()
#         ext = filename.lower().rsplit('.', 1)[-1]

#         if ext in {'doc', 'docx', 'docs'}:
#             try:
#                 extracted_text = process_document(content)
#             except Exception:
#                 extracted_text = "Failed to process Word document."
#         elif ext == 'pdf':
#             try:
#                 extracted_text = process_pdf(content)
#             except Exception:
#                 extracted_text = "Failed to process PDF document."
#         else:
#             extracted_text = "Unsupported file type"

#         if extracted_text and extracted_text.strip():
#             extracted_contents[filename] = extracted_text

#         tasks[task_id]["files"][filename] = {"status": "queued", "done": False}
#         asyncio.create_task(async_save_file(task_id, filename, content))

#     return jsonify({
#         "task_id": task_id,
#         "files": tasks[task_id]["files"],
#         "status": tasks[task_id]["status"]
#     })


# @app.route("/api/get_status/<task_id>")
# async def get_status(task_id):
#     if task_id not in tasks:
#         return jsonify({"error": "Invalid task ID"}), 404
#     return jsonify({
#         "task_id": task_id,
#         "files": tasks[task_id]["files"],
#         "status": tasks[task_id]["status"]
#     })


# async def call_openai_api(payload):
#     loop = asyncio.get_event_loop()
#     return await loop.run_in_executor(None, functools.partial(openai.ChatCompletion.create, **payload))


# @app.route("/submit-text", methods=["POST"])
# async def submit_text():
#     data = await request.get_json()
#     message = data.get("message")

#     try:
#         response = await call_openai_api({
#             "model": "gpt-4o-mini",
#             "messages": [
#                 {"role": "system", "content": "You are a helpful assistant. Use the following context to answer the user's question:"},
#                 {"role": "user", "content": f"Context:\n{extracted_contents}\n\nQuestion: {message}"}
#             ]
#         })
#         ai_reply = response['choices'][0]['message']['content']

#         return jsonify({
#             "status": "success",
#             "message": message,
#             "response": ai_reply
#         })

#     except Exception as e:
#         return jsonify({"status": "error", "error": str(e)}), 500


# if __name__ == "__main__":
#     app.run(debug=True)


from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import os, uuid, threading, time
from io import BytesIO
from docx import Document
import PyPDF2
import openai

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

tasks = {}
extracted_contents = {}

openai.api_key = os.getenv('OPENAI_API_KEY')  # Load from environment variable

# ======= Helper Functions =======

def process_document(file_contents):
    text = ''
    doc = Document(BytesIO(file_contents))
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text

def process_pdf(file_contents):
    text = ''
    reader = PyPDF2.PdfReader(BytesIO(file_contents))
    for page in reader.pages:
        text += page.extract_text()
    return text

def save_file(task_id, filename, content):
    tasks[task_id]["files"][filename] = {"status": "processing", "done": False}
    time.sleep(10)  # Simulate heavy processing
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    with open(filepath, "wb") as f:
        f.write(content)
    tasks[task_id]["files"][filename] = {"status": "completed", "done": True}

    if all(file["done"] for file in tasks[task_id]["files"].values()):
        tasks[task_id]["status"] = "completed"

# ======= Routes =======

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/upload", methods=["POST"])
def upload():
    files = request.files.getlist("files")
    if not files:
        return jsonify({"error": "No files uploaded"}), 400

    task_id = str(uuid.uuid4())
    tasks[task_id] = {"status": "processing", "files": {}}

    for f in files:
        filename = secure_filename(f.filename)
        content = f.read()

        ext = filename.lower().rsplit('.', 1)[-1]
        extracted_text = ""

        try:
            if ext in {'doc', 'docx', 'docs'}:
                extracted_text = process_document(content)
            elif ext == 'pdf':
                extracted_text = process_pdf(content)
            else:
                extracted_text = "Unsupported file type"
        except Exception as e:
            extracted_text = str(e)

        if extracted_text.strip():
            extracted_contents[filename] = extracted_text

        tasks[task_id]["files"][filename] = {"status": "queued", "done": False}
        thread = threading.Thread(target=save_file, args=(task_id, filename, content))
        thread.start()

    return jsonify({
        "task_id": task_id,
        "files": tasks[task_id]["files"],
        "status": tasks[task_id]["status"]
    })

@app.route("/api/get_status/<task_id>")
def get_status(task_id):
    if task_id not in tasks:
        return jsonify({"error": "Invalid task ID"}), 404
    return jsonify({
        "task_id": task_id,
        "files": tasks[task_id]["files"],
        "status": tasks[task_id]["status"]
    })

@app.route("/submit-text", methods=["POST"])
def submit_text():
    data = request.get_json()
    message = data.get("message")

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": f"Context:\n{extracted_contents}\n\nQuestion: {message}"}
            ]
        )
        ai_reply = response['choices'][0]['message']['content']
        return jsonify({
            "status": "success",
            "message": message,
            "response": ai_reply
        })
    except Exception as e:
        return jsonify({"status": "error", "error": str(e)}), 500

# ======= Run =======

if __name__ == "__main__":
    app.run(debug=True)

